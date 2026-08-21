#!/usr/bin/env python3
"""Build the customer-facing ERP product manual from the current Markdown docs."""

from __future__ import annotations

import argparse
import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BLUE = "2563EB"
CYAN = "14B8A6"
NAVY = "172033"
TEXT = "243047"
MUTED = "667085"
LIGHT = "F5F7FA"
LINE = "D9E2EC"
PALE_BLUE = "EFF6FF"
PALE_CYAN = "ECFDF9"

CHAPTERS = [
    ("登录、门店切换与权限", "docs/user-manual/01-local-login-and-permissions.md"),
    ("服务项目与价格管理", "docs/user-manual/02-service-items-and-pricing.md"),
    ("产品目录与标准价", "docs/user-manual/10-product-catalog-and-pricing.md"),
    ("设施接待与独立计时", "docs/user-manual/03-facility-reception-and-timing.md"),
    ("预约与员工排班", "docs/user-manual/18-appointments-and-employee-scheduling.md"),
    ("顾客档案与会员账户", "docs/user-manual/04-customers-and-membership.md"),
    ("会员储值与资金分账", "docs/user-manual/11-member-topups.md"),
    ("会员余额消费与手机号验证", "docs/user-manual/12-member-balance-payments.md"),
    ("次卡、积分与储值部分退款", "docs/user-manual/19-service-passes-points-and-partial-topup-refunds.md"),
    ("服务录单与金额确认", "docs/user-manual/05-service-order-and-price-confirmation.md"),
    ("收款、结算与收银交班", "docs/user-manual/06-payments-and-cashier-shifts.md"),
    ("消费退款与储值退款", "docs/user-manual/13-payment-refunds-and-topup-reversals.md"),
    ("商品销售与门店库存", "docs/user-manual/15-product-sales-and-inventory.md"),
    ("采购与供应链", None),
    ("产品图片与顾客服务档案", "docs/user-manual/16-product-images-and-service-records.md"),
    ("经营工作台与经营报表", "docs/user-manual/08-dashboard-and-operations-reports.md"),
    ("审计记录", "docs/user-manual/07-audit-events.md"),
    ("品牌与门店管理", "docs/user-manual/17-brand-and-store-management.md"),
    ("员工、账号与门店权限", "docs/user-manual/09-employees-accounts-and-permissions.md"),
    ("支付渠道配置与对账", "docs/user-manual/14-payment-channel-configuration.md"),
    ("商户入驻与平台管理", "docs/user-manual/20-platform-registration-security-and-administration.md"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, attrs in edges.items():
        tag = "start" if edge == "left" else "end" if edge == "right" else edge
        node = borders.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            borders.append(node)
        for key, value in attrs.items():
            node.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Arial Unicode MS", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_text_with_inline(paragraph, text: str, size=11, color=TEXT) -> None:
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            set_font(paragraph.add_run(text[pos:match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=size, color=NAVY, bold=True)
        elif token.startswith("`"):
            set_font(paragraph.add_run(token[1:-1]), name="Menlo", size=max(8.5, size - 1), color=BLUE)
        else:
            label = token[1:token.index("]")]
            set_font(paragraph.add_run(label), size=size, color=BLUE)
        pos = match.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]), size=size, color=color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def add_body(doc, text: str, *, bold=False, color=TEXT, after=6, align=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=after)
    if align is not None:
        p.alignment = align
    if bold:
        set_font(p.add_run(text), size=11, color=color, bold=True)
    else:
        add_text_with_inline(p, text, 11, color)
    return p


def add_heading(doc, text: str, level: int, page_break=False):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    if page_break:
        p.paragraph_format.page_break_before = True
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {1: NAVY, 2: BLUE, 3: NAVY}
    set_paragraph_spacing(p, before=10 if level > 1 else 0, after=7 if level > 1 else 12, line=1.05)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size=sizes[level], color=colors[level], bold=True)
    if level == 1:
        p._p.get_or_add_pPr().append(_bottom_rule(BLUE, 12, 8))
    return p


def _bottom_rule(color: str, size: int, space: int):
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    return p_bdr


def add_list_item(doc, text: str, ordered=False, level=0):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=3)
    p.paragraph_format.left_indent = Inches(0.22 + 0.22 * level)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    marker = f"{level + 1}." if ordered else "•"
    set_font(p.add_run(marker + "  "), size=10.5, color=BLUE, bold=True)
    add_text_with_inline(p, text, 11, TEXT)
    return p


def add_callout(doc, title: str, text: str, accent=CYAN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    cell = table.cell(0, 0)
    cell.width = Inches(6.35)
    set_cell_shading(cell, PALE_CYAN if accent == CYAN else PALE_BLUE)
    set_cell_margins(cell, top=130, start=170, bottom=130, end=170)
    set_cell_border(cell, left={"val": "single", "sz": 18, "color": accent})
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=3)
    set_font(p.add_run(title + "  "), size=10.5, color=accent, bold=True)
    add_text_with_inline(p, text, 10.5, TEXT)
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    width_count = max(len(r) for r in rows)
    clean = [r + [""] * (width_count - len(r)) for r in rows]
    table = doc.add_table(rows=len(clean), cols=width_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    total = 6.35
    if width_count == 2:
        ratios = [0.32, 0.68]
    elif width_count == 3:
        ratios = [0.20, 0.30, 0.50]
    elif width_count == 4:
        ratios = [0.17, 0.16, 0.28, 0.39]
    else:
        ratios = [1 / width_count] * width_count
    for row_index, row in enumerate(clean):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.width = Inches(total * ratios[col_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, "E8EEF5")
            elif row_index % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.12)
            add_text_with_inline(p, value.strip(), 9.2 if width_count >= 4 else 9.7, NAVY if row_index == 0 else TEXT)
            if row_index == 0:
                for run in p.runs:
                    run.bold = True
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def sanitize_text(text: str) -> str:
    replacements = {
        "V2 已开放": "已开放",
        "V2 已支持": "已支持",
        "当前版本": "当前系统",
        "本阶段": "当前范围",
        "第一版": "基础库存页面",
        "例如 `B202608200001`": "例如 `BYYYYMMDD0001`",
        "例如 B202608200001": "例如 BYYYYMMDD0001",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"V[12](?:\s*开发版|\s*当前开发基线)?", "当前系统", text)
    text = re.sub(r"2026[-/.年]0?8[-/.月](?:18|19|20|21|22)(?:日)?", "", text)
    text = re.sub(r"B2026\d{8}", "BYYYYMMDD0001", text)
    text = text.replace("`OWNER`", "负责人").replace("`STORE_MANAGER`", "店长")
    text = text.replace("`FRONT_DESK`", "前台").replace("`CASHIER`", "收银员").replace("`TECHNICIAN`", "服务员工")
    for duplicated, clean in {
        "系统负责人 负责人": "系统负责人",
        "门店店长 店长": "门店店长",
        "前台 前台": "前台",
        "收银员 收银员": "收银员",
        "服务员工 服务员工": "服务员工",
        "商户 负责人": "商户负责人",
    }.items():
        text = text.replace(duplicated, clean)
    return text.strip()


def should_drop_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    patterns = [
        r"^(适用版本|版本|更新日期|复核日期|日期)[:：]",
        r"V2026\d+__.*\.sql",
        r"\b(?:Git|commit|schema)\b",
        r"\d+\s*项(?:领域|集成|前端|自动化)?测试",
        r"P2-03",
        r"本地代码闭环",
    ]
    return any(re.search(pattern, stripped, re.I) for pattern in patterns)


def crop_image(source: Path, image_dir: Path) -> Path:
    target = image_dir / source.name
    with Image.open(source) as img:
        img = img.convert("RGB")
        crop_px = max(0, min(52, img.height // 14))
        if crop_px:
            img = img.crop((0, 0, img.width, img.height - crop_px))
        img.save(target, quality=92, optimize=True)
    return target


def add_image(doc, source: Path, alt: str, image_dir: Path, seen_images: set[Path]) -> None:
    source = source.resolve()
    if not source.exists() or source in seen_images:
        return
    seen_images.add(source)
    prepared = crop_image(source, image_dir)
    with Image.open(prepared) as img:
        ratio = img.height / img.width
    width = 6.25
    if ratio > 1.05:
        width = min(width, 5.15)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=2)
    shape = p.add_run().add_picture(str(prepared), width=Inches(width))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(caption, after=8)
    set_font(caption.add_run(f"图：{alt}"), size=9, color=MUTED)


def parse_markdown(
    doc,
    source: Path,
    image_dir: Path,
    seen_images: set[Path],
    *,
    skip_local_start=False,
    platform_customer_safe=False,
) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    table_rows: list[list[str]] = []
    in_code = False
    code_lines: list[str] = []
    skip_section = False

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, after=7, line=1.0)
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.right_indent = Inches(0.15)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT)
            p_pr.append(shd)
            set_font(p.add_run("\n".join(code_lines)), name="Menlo", size=8.5, color=NAVY)
            code_lines = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("# "):
            continue
        if skip_local_start and line.startswith("## 1. 启动本地系统"):
            skip_section = True
            continue
        if skip_section and line.startswith("## "):
            skip_section = False
        if skip_section:
            continue
        if platform_customer_safe and line.startswith("## 2. 初始化并登录平台管理员"):
            flush_table()
            add_heading(doc, "2. 初始化并登录平台管理员", 2)
            add_body(doc, "平台管理员由服务器管理员在系统首次启用时完成初始化。初始账号与一次性密码只通过受控渠道交付，不写入客户文档、共享聊天或普通业务页面。")
            add_body(doc, "平台管理员从独立的平台登录入口进入系统，首次登录必须立即修改初始密码。平台账号与商户员工账号相互独立，不能在商户的“员工与权限”页面创建。")
            skip_section = True
            continue
        if line.startswith("```"):
            flush_table()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if should_drop_line(line):
            continue
        if re.match(r"^\|?\s*:?-{3,}", line):
            continue
        if line.startswith("|") and line.endswith("|"):
            table_rows.append([sanitize_text(c.strip()) for c in line.strip("|").split("|")])
            continue
        flush_table()
        if not line.strip():
            continue
        image_match = re.match(r"!\[([^]]*)\]\(([^)]+)\)", line.strip())
        if image_match:
            add_image(doc, source.parent / image_match.group(2), image_match.group(1) or "功能界面", image_dir, seen_images)
            continue
        if line.startswith("### "):
            add_heading(doc, sanitize_text(re.sub(r"^###\s+", "", line)), 3)
        elif line.startswith("## "):
            heading = sanitize_text(re.sub(r"^##\s+", "", line))
            if skip_local_start:
                match = re.match(r"^(\d+)\.\s*(.*)$", heading)
                if match and int(match.group(1)) >= 2:
                    heading = f"{int(match.group(1)) - 1}. {match.group(2)}"
            add_heading(doc, heading, 2)
        elif re.match(r"^[-*]\s+", line):
            add_list_item(doc, sanitize_text(re.sub(r"^[-*]\s+", "", line)))
        elif re.match(r"^\d+[.)]\s+", line):
            add_list_item(doc, sanitize_text(re.sub(r"^\d+[.)]\s+", "", line)), ordered=True)
        elif line.startswith(">"):
            add_callout(doc, "提示", sanitize_text(line.lstrip("> ")))
        else:
            clean = sanitize_text(line)
            if clean:
                add_body(doc, clean)
    flush_table()
    flush_code()


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.30)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_font(run, size=9, color=MUTED)


def add_running_header_footer(section) -> None:
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, after=0)
    set_font(p.add_run("门店 ERP｜产品使用说明书"), size=8.5, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0)
    add_page_number(p)


def add_cover(doc: Document, logo: Path | None) -> None:
    for _ in range(3):
        doc.add_paragraph()
    if logo and logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p.add_run().add_picture(str(logo), width=Inches(0.78))
        shape._inline.docPr.set("descr", "门店 ERP 品牌标志")
        shape._inline.docPr.set("title", "门店 ERP 品牌标志")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=12)
    set_font(p.add_run("门店 ERP"), size=16, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=18, after=9, line=1.0)
    set_font(p.add_run("产品使用说明书"), size=32, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=26)
    set_font(p.add_run("连锁门店经营、会员、收银、库存与管理操作指南"), size=13, color=MUTED)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(1.45)
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, top=90, start=160, bottom=90, end=160)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0)
    set_font(p.add_run("客户使用版"), size=10.5, color=BLUE, bold=True)
    set_repeat_table_header(table.rows[0])
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("统一业务口径 · 权限边界 · 全流程操作"), size=9.5, color=MUTED)
    doc.add_page_break()


def add_part_page(doc: Document, part_no: str, title: str, subtitle: str) -> None:
    doc.add_page_break()
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=10)
    set_font(p.add_run(part_no), size=11, color=CYAN, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=12)
    set_font(p.add_run(title), size=26, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(subtitle), size=11, color=MUTED)


def add_contents(doc: Document) -> None:
    add_heading(doc, "使用说明", 1)
    add_callout(doc, "阅读建议", "新商户先阅读“系统概览”，再按实际岗位进入对应功能章节。页面是否可见取决于账号角色、门店范围和当前业务状态。")
    add_body(doc, "本手册面向品牌负责人、门店店长、前台、收银员、服务员工以及平台运营人员。它解释业务名词、操作顺序、字段规则、权限边界和常见状态，不包含开发过程、内部发布时间或工程实现记录。")
    add_heading(doc, "目录", 1)
    entries = [
        ("第一部分", "系统概览"),
        ("第 1 章", "产品定位与核心原则"),
        ("第 2 章", "业务闭环、菜单、角色与初次使用"),
        ("第二部分", "功能操作"),
    ]
    entries.extend((f"第 {index + 3} 章", title) for index, (title, _) in enumerate(CHAPTERS))
    entries.extend([("第三部分", "附录"), (f"第 {len(CHAPTERS) + 3} 章", "常用名词解释"), (f"第 {len(CHAPTERS) + 4} 章", "能力边界与问题处理")])
    add_table(doc, [["章节", "内容"]] + [[a, b] for a, b in entries])


def add_overview(doc: Document) -> None:
    add_heading(doc, "第 1 章  产品定位与核心原则", 1, page_break=True)
    add_body(doc, "门店 ERP 用于连接品牌总部、多个门店和不同岗位的日常经营。系统覆盖预约、接待、会员、收银、商品、库存、供应链、员工权限、审计和经营分析，并以品牌作为数据隔离边界。")
    principles = [
        ("设施计时与收费分离", "设施时长只记录占用；最终消费内容、服务时长和成交金额由有权限人员确认。"),
        ("服务与商品独立", "消费单可以只包含服务、只包含商品，也可以同时包含两者。"),
        ("价格统一发布", "负责人发布标准价；历史订单保留当时快照，后续改价不会重算旧单。"),
        ("品牌内共享、品牌间隔离", "同品牌多店共享顾客与会员权益，不同品牌的数据不能互查互用。"),
        ("资金、权益与库存使用流水", "已过账事实不直接覆盖；纠错通过退款、冲正或反向单据完成。"),
        ("前端隐藏不等于授权", "菜单、路由、按钮和后台接口都会校验角色、门店范围和业务状态。"),
        ("敏感信息按需展示", "姓名正常展示；列表手机号隐藏中间四位，完整手机号仍可精确查询。"),
        ("金额按分保存", "系统避免浮点误差，收银、退款和提成均按精确金额处理。"),
        ("业务编码自动生成", "数据库主键不对用户展示；服务、产品、员工、卡类、供应商和设施使用不可变业务编码。"),
    ]
    add_table(doc, [["原则", "业务含义"]] + [[a, b] for a, b in principles])
    add_callout(doc, "特别注意", "设施结束不会自动收费；商品退回库存不会自动退款；人工登记微信或支付宝也不会自动变成渠道已确认到账。", BLUE)

    add_heading(doc, "第 2 章  业务闭环、菜单、角色与初次使用", 1, page_break=True)
    add_heading(doc, "2.1 门店业务闭环", 2)
    flow = [
        "预约或顾客到店",
        "选择设施并开始记录占用",
        "结束服务或完成商品选购",
        "录入服务、商品、服务员工和成交价",
        "超出改价权限时由负责人审批",
        "确认应收金额并使用现金、会员或可信支付渠道结算",
        "库存、会员和支付流水同步入账",
        "提交交班，由不同人员复核",
        "通过审计记录和经营报表复盘",
    ]
    add_table(doc, [["顺序", "操作"]] + [[f"{i:02d}", item] for i, item in enumerate(flow, 1)])
    add_heading(doc, "2.2 左侧菜单速查", 2)
    menu = [
        ["经营工作台", "今日接待、设施、资金和待核对概况", "全部授权岗位"],
        ["设施接待", "现场资源占用、计时和清洁状态", "前台、店长、负责人"],
        ["预约与排班", "顾客预约与员工出勤计划", "前台、店长、负责人"],
        ["顾客与会员", "顾客档案、会员权益和服务档案", "授权服务与结算岗位"],
        ["服务录单与收银", "消费单、改价、收款、退款和交班", "收银员、店长、负责人"],
        ["商品库存", "余额、预占、流水和调整", "店长、负责人"],
        ["采购与供应链", "供应商、采购、批次、盘点和调拨", "店长、负责人"],
        ["服务项目", "维护可提供的服务和提成规则", "目录查看岗位、负责人"],
        ["产品目录", "维护可销售商品及可选图片", "目录查看岗位、负责人"],
        ["价格管理", "统一发布服务和商品标准价", "负责人"],
        ["经营报表", "收入、退款、支付、项目和员工绩效", "店长、负责人"],
        ["审计记录", "追溯关键操作", "授权管理岗位"],
        ["门店设施配置", "配置服务区、服务位和设施信息", "店长、负责人"],
        ["品牌与门店", "新增门店及维护组织生命周期", "负责人"],
        ["员工与权限", "员工档案、账号、角色和门店范围", "负责人"],
        ["支付渠道配置", "微信/支付宝凭据映射和对账", "负责人"],
    ]
    add_table(doc, [["菜单", "主要用途", "典型使用人"]] + menu)
    add_heading(doc, "2.3 角色与权限", 2)
    roles = [
        ["负责人", "品牌级最高权限、目录价格、组织员工、审批、审计与渠道配置", "全部业务和设置页面"],
        ["店长", "授权门店经营、设施、预约、顾客、收银、库存和报表", "不维护全局账号权限和支付密钥"],
        ["前台", "顾客预约、到店和设施现场操作", "工作台、设施、预约、顾客和只读目录"],
        ["收银员", "消费单、收款和本人收银班次", "工作台、顾客、收银和只读目录"],
        ["服务员工", "本人服务任务与记录的受限能力", "仅显示明确授予的入口"],
    ]
    add_table(doc, [["角色", "主要职责", "通常可见范围"]] + roles)
    add_body(doc, "同一员工可以兼任多个角色，最终权限取角色并集，但仍受门店范围、业务状态和职责分离规则约束。")
    add_heading(doc, "2.4 新商户初始化顺序", 2)
    setup = [
        "确认品牌资料并建立门店。",
        "创建店长、前台、收银员和服务员工，分配门店范围。",
        "建立服务项目和产品目录，按需上传产品图片。",
        "建立价格草稿，核对后由负责人发布。",
        "建立服务区、服务位和设施信息。",
        "配置会员卡类、储值、次卡和积分规则。",
        "需要商品库存时，录入期初、收货或采购入库。",
        "需要真实微信或支付宝时，先完成服务器凭据和商户验收。",
        "门店人员开始预约、接待、录单、收银、交班和复核。",
    ]
    for i, item in enumerate(setup, 1):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=5)
        set_font(p.add_run(f"{i:02d}"), size=10, color=CYAN, bold=True)
        set_font(p.add_run("   " + item), size=11, color=TEXT)


def add_supply_chain(doc: Document) -> None:
    add_heading(doc, "1. 功能范围", 2)
    add_body(doc, "“采购与供应链”集中处理供应商、采购入库、批次效期、库存盘点和跨店调拨。它与基础“商品库存”页面共同构成商品从入库、销售、盘点到跨店转移的追踪链路。")
    features = [
        ["供应商", "自动生成供应商编码；维护名称、联系人、联系电话、结算条款和启用状态"],
        ["采购入库", "记录供应商、门店、外部票号、说明、产品、数量、单位成本、批次号和可选有效期"],
        ["批次效期", "入库形成批次；出库优先使用最早到期批次，无有效期批次最后使用"],
        ["库存盘点", "提交时冻结账面数量，审批按实盘数与冻结数的差额过账"],
        ["跨店调拨", "按待出库、在途、已收货、已取消管理；出库和收货分别影响两店库存"],
    ]
    add_table(doc, [["功能", "用途"]] + features)
    add_heading(doc, "2. 权限与职责分离", 2)
    permissions = [
        ["查看供应商", "允许", "允许"],
        ["新增、编辑、停用供应商", "允许", "不允许"],
        ["查看及过账采购入库和采购成本", "允许", "不允许"],
        ["查看批次、数量和效期", "允许", "允许"],
        ["发起盘点", "允许", "限授权门店"],
        ["审批盘点", "允许，但不能审批本人申请", "不允许"],
        ["取消待审批盘点", "允许", "限授权门店"],
        ["新建、出库、收货和取消调拨", "允许", "不允许"],
    ]
    add_table(doc, [["操作", "负责人", "店长"]] + permissions)
    add_heading(doc, "3. 操作入口", 2)
    add_body(doc, "页面包含“批次效期、库存盘点、跨店调拨、供应商、采购入库”五个页签。采购入库、盘点和调拨均支持多产品行。")
    add_heading(doc, "4. 关键规则", 2)
    for text in [
        "采购入库、调拨出库、调拨收货和盘点审批属于正式过账动作，提交前应核对门店、产品、数量、批次和成本。",
        "盘点申请人与审批人必须是不同账号；盘点期间正常销售、退货或入库不会改变已经冻结的账面基准。",
        "调拨在确认出库时扣减调出店库存，在确认收货时增加调入店库存，并保留原批次、效期和单位成本。",
        "已过账记录不能直接编辑或物理删除；纠错必须通过新的反向业务单据保留完整历史。",
    ]:
        add_list_item(doc, text)
    add_heading(doc, "5. 当前边界", 2)
    add_body(doc, "当前供应链不等同于完整采购财务系统。采购订单、采购退货、供应商应付与付款、多仓库库位、税额、加权成本结转和会计凭证尚未开放。")


def add_terms(doc: Document) -> None:
    add_heading(doc, f"第 {len(CHAPTERS) + 3} 章  常用名词解释", 1, page_break=True)
    groups = [
        ("门店、接待与排班", [
            ["品牌/商户", "共享顾客和会员数据的经营主体，也是跨品牌隔离边界", "不等于一家具体门店"],
            ["门店上下文", "顶部“当前门店”选择的本次操作范围", "不会扩大账号门店权限"],
            ["服务区", "门店内组织服务位的逻辑分组", "不要求写死行业名称"],
            ["服务位/设施", "可被占用和计时的房间、床位或仪器", "不等于服务项目，也不自动定价"],
            ["设施占用时长", "服务位实际被记录占用的时间", "不等于收费时长或金额"],
            ["员工排班", "员工计划上班和下班的时间段", "不等于收银班次"],
        ]),
        ("目录、价格与订单", [
            ["标准价", "负责人发布的基准销售价格", "不是设施参考单价"],
            ["成交价", "本次订单实际向顾客收取的单价", "偏离标准价时可能需要审批"],
            ["价格版本", "某次发布时全部有效价格的完整快照", "已发布版本不覆盖修改"],
            ["待录单接待", "已结束接待但尚未形成消费单", "设施记录号仅用于追溯"],
            ["待支付", "应收已经确认，等待支付", "不等于已经收款"],
            ["作废", "终止未完成订单并保留历史", "不是物理删除"],
        ]),
        ("收银、会员与库存", [
            ["收银班次", "从开班到提交交班的一段收银责任区间", "不是员工排班"],
            ["交班复核", "另一名人员核对理论现金、实点现金和差额", "原收银员不能自我复核"],
            ["外部待核对", "人工登记或尚未被渠道确认的金额", "不能视为渠道到账证明"],
            ["储值本金", "顾客真实支付后进入会员账户的钱", "消费时优先扣除"],
            ["赠送奖励", "门店随储值额外赠送的权益", "不能直接退现金"],
            ["销售预占", "待支付订单暂时锁住的商品数量", "作废释放，结算转为出库"],
            ["过账", "让业务单据正式影响库存或账务", "过账后用反向单据纠错"],
            ["审计记录", "关键操作的不可变追溯事件", "不等于渠道回单或会计凭证"],
        ]),
    ]
    for heading, rows in groups:
        add_heading(doc, heading, 2)
        add_table(doc, [["名词", "通俗解释", "需要注意"]] + rows)


def add_boundaries(doc: Document) -> None:
    add_heading(doc, f"第 {len(CHAPTERS) + 4} 章  能力边界与问题处理", 1, page_break=True)
    add_heading(doc, "1. 需要外部条件后才能使用的能力", 2)
    add_body(doc, "微信支付和支付宝的交易、验签、查单、退款及账单对账，需要配置真实商户资料、证书、密钥和回调域名，并完成渠道验收。未完成验收前，系统内的人工收款记录不能作为渠道到账证明。")
    add_heading(doc, "2. 尚未纳入当前范围的能力", 2)
    for item in [
        "完整采购订单、采购退货、供应商应付付款和会计凭证。",
        "员工工资发放；当前提成仅作为收益核算和报表数据。",
        "复杂退卡作价、多次充值跨单追溯和异常非原路退款。",
        "自定义角色、临时授权和可视化权限设计器。",
        "平台模拟登录商户、跨品牌业务查看和物理删除商户。",
    ]:
        add_list_item(doc, item)
    add_heading(doc, "3. 常见问题处理", 2)
    faq = [
        ["菜单或按钮看不到", "确认账号角色、所属门店和对象状态；前端隐藏后，后台接口仍会校验权限"],
        ["手机号列表已隐藏", "这是视觉脱敏；输入完整手机号仍可精确查询，查看完整号码需说明用途并留痕"],
        ["设施已结束但仍未收款", "到“服务录单与收银”创建或打开消费单，确认金额并完成支付"],
        ["人工微信/支付宝仍待核对", "这是正常状态；需根据交易参考号和外部账单完成核对"],
        ["已过账记录录错", "不要删除或覆盖原事实；从原单发起退款、冲正或反向库存单据"],
        ["交班无法复核", "复核必须由另一账号完成；店长或负责人可按门店权限承担复核职责"],
        ["保存提示数据已变化", "表示其他人先完成了修改；刷新页面后重新核对再提交"],
    ]
    add_table(doc, [["问题", "处理建议"]] + faq)
    add_callout(doc, "安全原则", "遇到金额、会员权益、库存或跨门店数据异常时，先停止重复操作并保留页面提示、业务单号和追踪号，再由授权管理人员查询审计记录。", BLUE)


def scrub_docx_metadata(path: Path) -> None:
    temp = path.with_suffix(".scrubbed.docx")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "docProps/core.xml":
                root = ET.fromstring(data)
                for child in list(root):
                    local = child.tag.rsplit("}", 1)[-1]
                    if local in {"created", "modified", "lastPrinted", "revision"}:
                        root.remove(child)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            zout.writestr(item, data)
    temp.replace(path)


def document_text(path: Path) -> str:
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            texts.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        texts.extend(p.text for p in section.header.paragraphs)
        texts.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(texts)


def assert_customer_safe(path: Path) -> None:
    text = document_text(path)
    forbidden = [
        r"2026[-/.年]0?8[-/.月](?:18|19|20|21|22)",
        r"8月(?:18|19|20|21|22)日?",
        r"更新日期|复核日期|完成日期",
        r"开发版|当前开发基线",
        r"(?<!/)V[12]\b|V2026",
        r"P2-03|项测试|commit|提交哈希|迁移脚本",
    ]
    hits = [pattern for pattern in forbidden if re.search(pattern, text)]
    if hits:
        raise SystemExit(f"Customer manual contains forbidden internal metadata: {hits}")


def build(output: Path, logo: Path | None) -> None:
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_running_header_footer(doc.sections[0])
    core = doc.core_properties
    core.title = "门店 ERP 产品使用说明书"
    core.subject = "连锁门店经营管理系统客户使用指南"
    core.author = "门店 ERP"
    core.keywords = "门店, ERP, 使用说明书"
    core.comments = "客户使用版"

    with tempfile.TemporaryDirectory(prefix="erp-manual-images-") as tmp:
        image_dir = Path(tmp)
        seen_images: set[Path] = set()
        add_cover(doc, logo)
        add_contents(doc)
        add_part_page(doc, "PART I", "系统概览", "先理解业务边界，再开始实际操作")
        add_overview(doc)
        add_part_page(doc, "PART II", "功能操作", "按岗位和业务场景查阅对应章节")
        for offset, (title, rel_path) in enumerate(CHAPTERS, start=3):
            add_heading(doc, f"第 {offset} 章  {title}", 1, page_break=True)
            if rel_path is None:
                add_supply_chain(doc)
            else:
                parse_markdown(
                    doc,
                    ROOT / rel_path,
                    image_dir,
                    seen_images,
                    skip_local_start=rel_path.endswith("01-local-login-and-permissions.md"),
                    platform_customer_safe=rel_path.endswith("20-platform-registration-security-and-administration.md"),
                )
        add_part_page(doc, "PART III", "附录", "名词解释、能力边界与问题处理")
        add_terms(doc)
        add_boundaries(doc)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)
    scrub_docx_metadata(output)
    assert_customer_safe(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=ROOT / "门店ERP产品使用说明书.docx")
    parser.add_argument("--logo", type=Path)
    args = parser.parse_args()
    build(args.output.resolve(), args.logo.resolve() if args.logo else None)
    print(args.output.resolve())


if __name__ == "__main__":
    main()
